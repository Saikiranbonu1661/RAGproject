"""
Document Processing Module

Advanced document processing with:
- Unstructured data loading
- Regex-based text cleaning
- Tiktoken-based chunking
- Custom metadata handling
"""

import os
import re
import logging
import datetime
from typing import List, IO, Dict, Any
from abc import ABC, abstractmethod
import warnings
warnings.filterwarnings("ignore")

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

# LangChain
from langchain.schema import Document
from langchain_community.document_loaders import UnstructuredFileIOLoader

# Local imports
from .text_splitter import TextSplitter

# Setup logging
logger = logging.getLogger(__name__)

# Regex patterns for text cleaning
pattern1 = re.compile(r'\s+')  # Multiple whitespaces
pattern2 = re.compile(r'\n+')  # Multiple newlines
pattern3 = re.compile(r'[^\w\s.,!?;:()\[\]{}\-\'\"]+')  # Special characters


class DataLoaders(ABC):
    """Abstract base class for data loaders."""
    
    @abstractmethod
    def load_data(self, data):
        raise NotImplementedError
    
    @staticmethod
    def update_metadata_in_documents(metadata: Dict[str, Any], documents: List[Document]) -> List[Document]:
        """Update metadata in all documents."""
        for document in documents:
            document.metadata.update(metadata)
        return documents


class FileIoDataLoader(DataLoaders):
    """Advanced file data loader with text cleaning and chunking."""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100, 
                 separators: List[str] = None):
        """
        Initialize the FileIoDataLoader.
        
        Args:
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
            separators: List of separators for recursive chunking
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", " ", ""]
        logger.debug("FileIoDataLoader initialized")

    def load_data(self, io_data: IO) -> List[Document]:
        """
        Load and clean data from IO stream.
        
        Args:
            io_data: IO stream of file data
            
        Returns:
            List of cleaned and chunked documents
        """
        # Load documents using Unstructured
        loader = UnstructuredFileIOLoader(io_data)
        documents = loader.load()
        
        # Clean text content with regex patterns
        for document in documents:
            # Replace multiple whitespaces with single space
            document.page_content = re.sub(pattern1, ' ', document.page_content)
            # Replace multiple newlines with single newline
            document.page_content = re.sub(pattern2, '\n', document.page_content)
            # Remove unwanted special characters
            document.page_content = re.sub(pattern3, '', document.page_content)
        
        # Split documents using TextSplitter
        text_splitter = TextSplitter()
        text_splitter.chunk_size = self.chunk_size
        text_splitter.chunk_overlap = self.chunk_overlap
        text_splitter.separators = self.separators
        final_documents = text_splitter.split_documents(documents)
        
        return final_documents

    def scrap_and_create_documents_for_file_data(
        self, 
        bytes_data: IO, 
        file_path: str,
        file_name: str
    ) -> List[Document]:
        """
        Process file data with metadata.
        
        Args:
            bytes_data: IO stream of file data
            file_path: Path to the source file
            file_name: Name of the file
            
        Returns:
            List of processed documents with metadata
        """
        try:
            # Create metadata
            metadata = {
                'content_type': 'file',
                'file_name': file_name,
                'timestamp': datetime.datetime.now().isoformat(),
                'source': file_path,
            }
            
            # Load and process documents
            documents = self.load_data(bytes_data)
            
            # Update metadata in all documents
            final_documents = self.update_metadata_in_documents(metadata, documents)
            
            logger.debug(f"Processed {len(final_documents)} chunks from {file_name}")
            return final_documents
            
        except Exception as e:
            logger.error(f"Error processing file {file_name}: {e}")
            raise e


class DocumentProcessor:
    """
    Enhanced document processor with advanced chunking and cleaning.
    
    Maintains backward compatibility while adding support for:
    - Unstructured file loading
    - Advanced text cleaning
    - Tiktoken-based chunking
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        """Initialize the document processor."""
        self.supported_formats = ['.pdf', '.txt', '.docx', '.md']
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.file_loader = FileIoDataLoader(chunk_size, chunk_overlap)
        logger.debug("DocumentProcessor initialized")
    
    def load_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF files.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                logger.debug(f"Extracted text from PDF: {file_path}")
                return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX files.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            logger.debug(f"Extracted text from DOCX: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """
        Load text from TXT files.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            File content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                logger.debug(f"Loaded TXT file: {file_path}")
                return content
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def load_markdown(self, file_path: str) -> str:
        """
        Load and convert markdown to text.
        
        Args:
            file_path: Path to the Markdown file
            
        Returns:
            Converted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                text = soup.get_text()
                logger.debug(f"Converted Markdown file: {file_path}")
                return text
        except Exception as e:
            logger.error(f"Error reading Markdown {file_path}: {e}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """
        Load document based on file extension.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Document content as string
        """
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return ""
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.load_pdf(file_path)
        elif file_ext == '.docx':
            return self.load_docx(file_path)
        elif file_ext == '.txt':
            return self.load_txt(file_path)
        elif file_ext == '.md':
            return self.load_markdown(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return ""
    
    def load_document_advanced(self, file_path: str) -> List[Document]:
        """
        Load and process document using advanced FileIoDataLoader.
        
        Uses:
        - UnstructuredFileIOLoader for loading
        - Regex-based text cleaning
        - Tiktoken-based chunking
        - Rich metadata
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of processed and chunked Document objects
        """
        try:
            with open(file_path, 'rb') as file:
                filename = os.path.basename(file_path)
                data_id = os.path.splitext(filename)[0]
                
                documents = self.file_loader.scrap_and_create_documents_for_file_data(
                    bytes_data=file,
                    file_path=file_path,
                    file_name=filename,
                    data_id=data_id,
                    workflow_id='rag_qa'
                )
                
                logger.debug(f"Processed {len(documents)} chunks from {filename}")
                return documents
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return []
    
    def load_documents_from_directory(self, directory_path: str, use_advanced: bool = True) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to directory containing documents
            use_advanced: Use advanced FileIoDataLoader (default: True)
            
        Returns:
            List of LangChain Document objects
        """
        if not os.path.exists(directory_path):
            logger.error(f"Directory does not exist: {directory_path}")
            return []
            
        documents = []
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(filename)[1].lower()
                
                if file_ext in self.supported_formats:
                    logger.debug(f"Processing document: {filename}")
                    
                    if use_advanced:
                        # Use advanced processing with tiktoken chunking
                        docs = self.load_document_advanced(file_path)
                        documents.extend(docs)
                    else:
                        # Use simple loading (backward compatibility)
                        content = self.load_document(file_path)
                        
                        if content.strip():
                            documents.append(Document(
                                page_content=content,
                                metadata={
                                    "source": filename,
                                    "path": file_path,
                                    "type": file_ext[1:]  # Remove the dot
                                }
                            ))
                        else:
                            logger.warning(f"No content extracted from: {filename}")
                else:
                    logger.info(f"Skipping unsupported file: {filename}")
        
        logger.info(f"Loaded {len(documents)} document chunks from {directory_path}")
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return self.supported_formats.copy()
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate if a file can be processed.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file can be processed, False otherwise
        """
        if not os.path.exists(file_path):
            return False
            
        if not os.path.isfile(file_path):
            return False
            
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_formats 